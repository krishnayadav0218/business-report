"""
test_ingest.py
Covers src/ingest.py: reading CSV/Excel files, handling messy real-world
formats (title rows, multiple sheets), and clear errors for unusable files.
"""

import pytest
from src import ingest


def test_reads_clean_csv(sample_csv_path):
    df = ingest.from_any_file(sample_csv_path)
    assert list(df.columns) == ["Date", "Region", "Salesperson", "Product", "Target", "Collection", "Orders", "Cost"]
    assert len(df) == 4


def test_handles_title_row_and_blank_row(messy_excel_path):
    """A company name + blank row before the real header shouldn't produce 'Unnamed' columns."""
    df = ingest.from_any_file(messy_excel_path)
    assert not any(str(c).lower().startswith("unnamed") for c in df.columns)
    assert "Sale Date" in df.columns
    assert len(df) == 2


def test_picks_the_sheet_with_real_data(multisheet_excel_path):
    """When sheet 1 is a cover page, the loader should find the sheet with actual data."""
    df = ingest.from_any_file(multisheet_excel_path)
    assert "Date" in df.columns
    assert "Collection" in df.columns
    assert len(df) == 2


def test_unsupported_extension_raises_clear_error(tmp_path):
    bad_file = tmp_path / "data.txt"
    bad_file.write_text("not a real data file")
    with pytest.raises(ValueError, match="Unsupported file type"):
        ingest.from_any_file(str(bad_file))


def test_old_doc_format_raises_helpful_error(tmp_path):
    fake_doc = tmp_path / "report.doc"
    fake_doc.write_text("legacy binary format placeholder")
    with pytest.raises(ValueError, match="\\.docx"):
        ingest.from_any_file(str(fake_doc))


def test_merged_two_row_header(tmp_path):
    """A top row with a merged label + a sub-header row below it should combine correctly."""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.append(["Date", "Region", "Sales Figures", None, None])
    ws.append([None, None, "Target", "Collection", "Orders"])
    ws.merge_cells("C1:E1")
    ws.append(["01-06-2026", "Mumbai", 50000, 48000, 12])
    path = tmp_path / "merged.xlsx"
    wb.save(path)

    df = ingest.from_any_file(str(path))
    assert "Target" in df.columns
    assert "Collection" in df.columns
    assert "Orders" in df.columns


def test_single_form_style_excel_is_transposed_to_one_row(tmp_path):
    """'Label: Value' rows (a form, not a table) should become one proper data row."""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.append(["Region", "West"])
    ws.append(["Date", "01-06-2026"])
    ws.append(["Collection", 50000])
    ws.append(["Target", 55000])
    path = tmp_path / "form.xlsx"
    wb.save(path)

    df = ingest.from_any_file(str(path))
    assert len(df) == 1
    assert set(["Region", "Date", "Collection", "Target"]).issubset(df.columns)
    assert df.iloc[0]["Collection"] == 50000


def test_multiple_form_blocks_become_multiple_rows(tmp_path):
    """Several Label:Value blocks separated by blank rows -- e.g. one per
    transaction -- should each become their own row in the final table."""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    for region, date_, collection in [("West", "01-06-2026", 50000), ("North", "02-06-2026", 60000), ("East", "03-06-2026", 45000)]:
        ws.append(["Region", region])
        ws.append(["Date", date_])
        ws.append(["Collection", collection])
        ws.append([])
    path = tmp_path / "blocks.xlsx"
    wb.save(path)

    df = ingest.from_any_file(str(path))
    assert len(df) == 3
    assert set(df["Region"]) == {"West", "North", "East"}


def test_pdf_with_label_value_lines_instead_of_a_table(tmp_path):
    """A PDF that's just paragraph-style 'Label: Value' lines (no ruled table)
    should still yield usable data instead of failing outright."""
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet

    path = tmp_path / "summary.pdf"
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    doc.build([
        Paragraph("Monthly Business Summary", styles["Title"]),
        Paragraph("Region: West", styles["Normal"]),
        Paragraph("Date: 01-06-2026", styles["Normal"]),
        Paragraph("Collection: 50000", styles["Normal"]),
    ])

    df = ingest.from_any_file(str(path))
    assert len(df) == 1
    assert "Collection" in df.columns


def test_docx_with_label_value_lines_instead_of_a_table(tmp_path):
    import docx as docx_lib
    path = tmp_path / "summary.docx"
    d = docx_lib.Document()
    d.add_paragraph("Region: West")
    d.add_paragraph("Date: 01-06-2026")
    d.add_paragraph("Collection: 50000")
    d.save(path)

    df = ingest.from_any_file(str(path))
    assert len(df) == 1
    assert "Collection" in df.columns


def test_real_table_still_preferred_over_key_value_parsing(sample_csv_path):
    """A genuine multi-row/multi-column table should never get accidentally
    treated as a key-value form -- regression guard for the new detection."""
    df = ingest.from_any_file(sample_csv_path)
    assert len(df) == 4  # the real row count, not collapsed to 1
