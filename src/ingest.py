"""
ingest.py
Data Ingestion Layer -- pulls raw data from whichever source is configured.

File types handled by from_any_file() (used by the web app / desktop app / watch folder):
    .csv           -> plain CSV
    .xlsx / .xls   -> Excel (scans ALL sheets, picks the one with the real data)
    .pdf           -> PDF containing a table OR "Label: Value" style lines
    .docx          -> Word document containing a table OR "Label: Value" style lines

Other sources (set DATA_SOURCE in .env, used by src/main.py for scheduled runs):
    gsheet  -> Google Sheets (via gspread)                (needs service_account.json)
    sql     -> PostgreSQL / MySQL                         (needs DB creds in .env)

The rest of the pipeline only cares about getting back a pandas DataFrame,
so swapping sources never touches process.py, charts.py, or report_builder.py.

ROBUSTNESS -- every file type is read "headerless" first and passed through
_resolve_table(), which protects against, so more than just clean tables work:

1. HEADER ROW NOT ON ROW 1
   Company files often have a title/company name or blank rows above the real
   header. Scored by "looks like a row of text labels", not just "has
   content" -- so a numeric data row is never mistaken for the header just
   because it's fully filled in.

2. TWO-ROW / MERGED HEADERS
   e.g. a top row "Sales Figures" spanning several merged columns, with the
   real sub-headers ("Target", "Collection", "Orders") on the row below. The
   two rows are combined automatically, preferring the more specific
   (bottom) label per column when both exist.

3. MULTIPLE SHEETS IN ONE EXCEL FILE
   If the real data isn't on the first sheet (e.g. sheet 1 is a cover page),
   every sheet is scanned and the one that yields the most usable rows +
   columns is used automatically.

4. BLANK ROWS/COLUMNS MID-FILE
   Fully-empty rows and columns are dropped after reading, wherever they are.

5. "FORM" / VERTICAL LAYOUTS (not a row-and-column table at all)
   Some Excel/CSV files, or PDF/Word documents, list data as
       Region: West
       Collection: 48000
       Date: 01/06/2026
   instead of a table. This is auto-detected and pivoted into a proper row --
   if the sheet has several such blocks separated by blank rows (e.g. one
   block per transaction), each block becomes its own row. PDFs/Word docs
   without any real table get the same "Label: Value" line-matching applied
   to their text as a last resort before giving up.

   Note: a single free-floating document (one invoice, one letter) can only
   ever produce ONE row of data this way, so charts that need multiple rows
   (trends, regional comparison) won't have much to show -- but the summary
   numbers will still come through instead of a hard failure.
"""

import os
import re
import pandas as pd

MAX_HEADER_SCAN_ROWS = 20
_KV_LINE_RE = re.compile(r"^(.{2,40}?)\s*[:\-]\s+(.+)$")


def _is_textlike(val) -> bool:
    """True if a cell value reads as a text label rather than a number."""
    text = str(val).strip()
    if not text:
        return False
    try:
        float(text.replace(",", "").replace("\u20b9", "").replace("$", ""))
        return False
    except ValueError:
        return True


def _row_header_score(row: pd.Series) -> float:
    """
    Higher score = more likely this row is a header row of column labels.
    Rewards cells that are filled AND text-like (headers are almost always
    short text labels); a fully-numeric row (a data row) scores much lower
    even if every cell is filled.
    """
    score = 0.0
    for val in row:
        if pd.isna(val):
            continue
        text = str(val).strip()
        if not text:
            continue
        score += 1.0 if _is_textlike(val) else 0.3
    return score


def _detect_header_row(raw_df: pd.DataFrame) -> int:
    """Scan the first N rows and return the index of the best-scoring header row."""
    scan_limit = min(MAX_HEADER_SCAN_ROWS, len(raw_df))
    best_idx, best_score = 0, -1.0
    for i in range(scan_limit):
        score = _row_header_score(raw_df.iloc[i])
        if score > best_score:
            best_score = score
            best_idx = i
    return best_idx


def _clean_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Strip whitespace/newlines from headers, drop fully-empty rows/columns."""
    df = df.dropna(axis=1, how="all")
    df = df.dropna(axis=0, how="all")
    df.columns = [str(c).strip() for c in df.columns]
    return df


# ---------------------------------------------------------------------------
# "Form" / vertical Label:Value layout support
# ---------------------------------------------------------------------------

def _looks_like_key_value_block(block: pd.DataFrame) -> bool:
    """
    True if a block of rows reads like a vertical form (label in column 0,
    value in column 1) rather than a real multi-column table: few columns,
    column 0 is mostly text, and labels are mostly unique (not repeating
    category names the way a real table's first column often is).
    """
    if block.shape[1] > 3 or block.shape[1] < 2:
        return False
    if len(block) < 2:
        return False
    col0 = block.iloc[:, 0].dropna()
    if len(col0) == 0:
        return False
    textlike_ratio = sum(_is_textlike(v) for v in col0) / len(col0)
    if textlike_ratio < 0.7:
        return False
    labels = col0.astype(str).str.strip()
    return (labels.nunique() / len(labels)) > 0.8


def _key_value_block_to_row(block: pd.DataFrame) -> dict:
    """Turns a 'Label: Value' block into a single dict -- one row of a table."""
    row = {}
    for _, r in block.iterrows():
        label = r.iloc[0]
        value = r.iloc[1] if len(r) > 1 else None
        if pd.isna(label) or str(label).strip() == "":
            continue
        row[str(label).strip()] = value
    return row


def _split_into_blocks(raw: pd.DataFrame) -> list:
    """Splits a sheet into consecutive non-blank-row groups, separated by fully-blank rows."""
    blocks, current = [], []
    for _, row in raw.iterrows():
        if row.isna().all():
            if current:
                blocks.append(pd.DataFrame(current))
                current = []
        else:
            current.append(row)
    if current:
        blocks.append(pd.DataFrame(current))
    return blocks


def _try_key_value_layout(raw: pd.DataFrame):
    """
    Returns a resolved DataFrame if `raw` looks like one or more Label:Value
    forms (rather than a normal table), or None if it doesn't apply -- in
    which case the caller falls through to normal header-row detection.
    """
    blocks = _split_into_blocks(raw)
    real_blocks = [b for b in blocks if len(b) >= 2]

    if len(real_blocks) > 1 and all(_looks_like_key_value_block(b) for b in real_blocks):
        rows = [_key_value_block_to_row(b) for b in real_blocks]
        return _clean_columns(pd.DataFrame(rows))

    if len(real_blocks) == 1 and _looks_like_key_value_block(raw):
        return _clean_columns(pd.DataFrame([_key_value_block_to_row(raw)]))

    return None


def _extract_key_value_pairs_from_text(text: str) -> dict:
    """Pulls 'Label: Value' or 'Label - Value' lines out of free-form text
    (used as a last resort for PDFs/Word docs with no real table)."""
    pairs = {}
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        m = _KV_LINE_RE.match(line)
        if m:
            label, value = m.group(1).strip(), m.group(2).strip()
            if label and value:
                pairs[label] = value
    return pairs


def _resolve_table(raw: pd.DataFrame) -> pd.DataFrame:
    """
    Given a raw (headerless) DataFrame -- straight from a CSV/Excel sheet/
    PDF or Word table -- detect where the real header is (one row, or two
    rows for a merged/multi-level header) and return the finished, properly
    labeled data table. Also detects vertical "Label: Value" form layouts
    and pivots them into a normal table first.
    """
    if raw is None or raw.empty:
        raise ValueError("This file/sheet has no data.")

    kv_result = _try_key_value_layout(raw)
    if kv_result is not None:
        return kv_result

    idx = _detect_header_row(raw)
    header_rows = [idx]

    # Check whether the next row is a genuine second header level (common with
    # merged cells): some columns blank in the detected header but filled with
    # text in the row right below it.
    if idx + 1 < len(raw):
        row0 = raw.iloc[idx]
        row1 = raw.iloc[idx + 1]
        score0 = _row_header_score(row0)
        score1 = _row_header_score(row1)
        complementary = sum(
            1 for h, n in zip(row0, row1)
            if (pd.isna(h) or str(h).strip() == "") and not pd.isna(n) and _is_textlike(n)
        )
        if complementary >= 1 and score0 > 0 and score1 >= score0 * 0.4:
            header_rows.append(idx + 1)

    if len(header_rows) == 2:
        top = raw.iloc[header_rows[0]].ffill()  # spread a merged top label across its columns
        bottom = raw.iloc[header_rows[1]]
        columns = []
        for t, b in zip(top, bottom):
            b_str = "" if pd.isna(b) else str(b).strip()
            t_str = "" if pd.isna(t) else str(t).strip()
            columns.append(b_str if b_str else t_str)  # prefer the more specific sub-header
        data_start = header_rows[1] + 1
    else:
        header = raw.iloc[header_rows[0]].ffill()
        columns = [("" if pd.isna(c) else str(c).strip()) for c in header]
        data_start = header_rows[0] + 1

    df = raw.iloc[data_start:].copy()
    df.columns = columns
    return _clean_columns(df)


def from_csv(path: str) -> pd.DataFrame:
    """Read raw data from a local CSV (or any URL pandas can read), auto-detecting the header."""
    raw = pd.read_csv(path, header=None)
    return _resolve_table(raw)


def _score_sheet(df: pd.DataFrame) -> tuple:
    """Score a candidate sheet by (usable column count, data row count) -- used to
    pick the right sheet when an Excel file has more than one."""
    if df is None or df.empty:
        return (0, 0)
    usable_cols = sum(1 for c in df.columns if c and not str(c).lower().startswith("unnamed"))
    return (usable_cols, len(df))


def _pick_best_excel_sheet(path: str, engine) -> pd.DataFrame:
    """Read every sheet in the workbook and keep whichever one looks like the real data table."""
    xls = pd.ExcelFile(path, engine=engine)
    best_df, best_score = None, (-1, -1)
    for sheet in xls.sheet_names:
        try:
            raw = pd.read_excel(xls, sheet_name=sheet, header=None)
            if raw.empty or len(raw) < 2:
                continue
            df = _resolve_table(raw)
        except Exception:
            continue
        score = _score_sheet(df)
        if score > best_score:
            best_score = score
            best_df = df
    if best_df is None:
        raise ValueError(
            "Couldn't find a usable data table in any sheet of this Excel file. "
            "Make sure at least one sheet has a row of column headers followed by data rows."
        )
    return best_df


def from_excel(path: str, sheet_name=None) -> pd.DataFrame:
    """
    Read raw data from a local .xlsx/.xls file. If sheet_name is given, reads
    only that sheet; otherwise scans every sheet and auto-picks the one with
    real tabular data. Auto-detects the header row(s) either way.
    """
    engine = "openpyxl" if str(path).lower().endswith("x") else None
    if sheet_name is not None:
        raw = pd.read_excel(path, sheet_name=sheet_name, header=None, engine=engine)
        return _resolve_table(raw)
    return _pick_best_excel_sheet(path, engine)


def from_pdf(path: str) -> pd.DataFrame:
    """
    Extract tabular data from a PDF. Scans every page for tables and picks the
    largest one -- or, if several pages share the same header row (a table that
    continues across pages), stitches them together into a single DataFrame.
    If no real table is found (or what looks like a "table" is actually just
    paragraph text pdfplumber misread), falls back to matching "Label: Value"
    style lines in the page text.
    """
    import pdfplumber

    all_tables = []  # list of (page_num, rows)
    all_text = ""
    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            # Try the default table-detection strategy first, then a looser
            # text-based strategy as a fallback for tables with faint/no ruling lines.
            tables = page.extract_tables()
            if not tables:
                tables = page.extract_tables(table_settings={
                    "vertical_strategy": "text", "horizontal_strategy": "text"
                })
            for table in tables:
                if table and len(table) >= 2:  # needs at least a header + 1 data row
                    all_tables.append((page_num, table))
            all_text += (page.extract_text() or "") + "\n"

    table_result = None
    if all_tables:
        # If multiple tables share an identical first row, they're likely the same
        # table continuing across pages -- stitch them together.
        first_header = tuple(str(c) for c in all_tables[0][1][0])
        matching = [rows for _, rows in all_tables if tuple(str(c) for c in rows[0]) == first_header]
        if len(matching) > 1:
            combined = [matching[0][0]]
            for rows in matching:
                combined.extend(rows[1:])
            candidate = pd.DataFrame(combined)
        else:
            _, biggest = max(all_tables, key=lambda t: len(t[1]))
            candidate = pd.DataFrame(biggest)

        # A "table" with only one column is almost always the text-based
        # fallback strategy misreading paragraph text, not a real table --
        # ignore it and let the Label:Value text parser handle this file instead.
        if candidate.shape[1] >= 2:
            table_result = _resolve_table(candidate)

    if table_result is not None:
        return table_result

    pairs = _extract_key_value_pairs_from_text(all_text)
    if len(pairs) >= 2:
        return _clean_columns(pd.DataFrame([pairs]))

    raise ValueError(
        "No table -- and no clear 'Label: Value' lines -- could be found in this PDF. "
        "This works best with a real data table, or a document with lines like "
        "'Region: West' / 'Collection: 48000'. If this is a scanned/image PDF, "
        "text can't be read from it without OCR. Try exporting the data as Excel/CSV instead."
    )


def from_docx(path: str) -> pd.DataFrame:
    """
    Extract tabular data from a Word document (.docx). Considers every table in
    the document and picks the one that looks most like real data (most rows x
    columns), not just the first one. .doc (old binary Word format) isn't
    supported -- save as .docx first.
    """
    import docx

    document = docx.Document(path)
    best_rows, best_score = None, (-1, -1)
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows and len(rows[0]) >= 2:  # ignore single-column "tables" (rare, but same guard as PDFs)
            score = (len(rows), len(rows[0]))
            if score > best_score:
                best_score = score
                best_rows = rows

    if best_rows is not None:
        return _resolve_table(pd.DataFrame(best_rows))

    all_text = "\n".join(p.text for p in document.paragraphs)
    pairs = _extract_key_value_pairs_from_text(all_text)
    if len(pairs) >= 2:
        return _clean_columns(pd.DataFrame([pairs]))

    raise ValueError(
        "No table -- and no clear 'Label: Value' lines -- could be found in this Word "
        "document. Either put your data in a real Word table (Insert > Table), or use "
        "lines like 'Region: West' / 'Collection: 48000'."
    )


def from_any_file(path: str) -> pd.DataFrame:
    """Auto-detect file type by extension -- used by the drag-and-drop tool and web app."""
    ext = os.path.splitext(path)[1].lower()
    if ext in (".xlsx", ".xls"):
        return from_excel(path)
    elif ext == ".csv":
        return from_csv(path)
    elif ext == ".pdf":
        return from_pdf(path)
    elif ext == ".docx":
        return from_docx(path)
    elif ext == ".doc":
        raise ValueError("Old .doc format isn't supported -- please save/export as .docx and try again.")
    else:
        raise ValueError(f"Unsupported file type: {ext} (use .xlsx, .xls, .csv, .pdf, or .docx)")


def from_google_sheets(sheet_url: str, worksheet_name: str = "Sheet1") -> pd.DataFrame:
    """
    Read a live Google Sheet.
    Needs: pip install gspread oauth2client
           a service_account.json (share the sheet with its client_email)
    """
    import gspread
    from oauth2client.service_account import ServiceAccountCredentials

    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    creds_path = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON", "service_account.json")
    creds = ServiceAccountCredentials.from_json_keyfile_name(creds_path, scope)
    client = gspread.authorize(creds)

    sheet = client.open_by_url(sheet_url).worksheet(worksheet_name)
    records = sheet.get_all_records()
    return pd.DataFrame(records)


def from_sql(query: str) -> pd.DataFrame:
    """
    Read from PostgreSQL or MySQL depending on DB_ENGINE in .env.
    Needs: pip install sqlalchemy psycopg2-binary   (postgres)
                       pymysql                       (mysql)
    """
    from sqlalchemy import create_engine

    engine_type = os.getenv("DB_ENGINE", "postgresql")  # or "mysql+pymysql"
    user = os.getenv("DB_USER")
    password = os.getenv("DB_PASSWORD")
    host = os.getenv("DB_HOST")
    port = os.getenv("DB_PORT")
    dbname = os.getenv("DB_NAME")

    conn_str = f"{engine_type}://{user}:{password}@{host}:{port}/{dbname}"
    engine = create_engine(conn_str)
    return pd.read_sql(query, engine)


def from_api(url: str, headers: dict | None = None) -> pd.DataFrame:
    """Pull JSON data from a third-party API (e.g. Stripe, GA) and flatten to a DataFrame."""
    import requests

    resp = requests.get(url, headers=headers or {}, timeout=30)
    resp.raise_for_status()
    return pd.json_normalize(resp.json())


def load_data() -> pd.DataFrame:
    """
    Single entry point used by main.py.
    Reads DATA_SOURCE from the environment and dispatches to the right loader.
    Defaults to the bundled sample CSV so the pipeline runs with zero config.
    """
    source = os.getenv("DATA_SOURCE", "csv")

    if source == "csv":
        path = os.getenv("CSV_PATH", "data/sample_sales_data.csv")
        return from_csv(path)
    elif source == "excel":
        path = os.getenv("EXCEL_PATH", "data/sample_sales_data.xlsx")
        return from_excel(path)
    elif source == "gsheet":
        return from_google_sheets(os.getenv("GOOGLE_SHEET_URL"))
    elif source == "sql":
        return from_sql(os.getenv("SQL_QUERY", "SELECT * FROM sales;"))
    elif source == "api":
        return from_api(os.getenv("API_URL"))
    else:
        raise ValueError(f"Unknown DATA_SOURCE: {source}")
